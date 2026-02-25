from typing import Dict, Any, List
from io import BytesIO
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

class WordExporter:
    """Word导出器"""

    def _setup_document(self) -> Document:
        """创建并设置文档基础配置"""
        document = Document()

        # 设置默认中文字体
        document.styles['Normal'].font.name = 'SimSun'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

        # 设置A4纸张
        section = document.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

        return document

    def _extract_content(self, content: Any) -> str:
        """从 content 中提取剧本文本"""
        full_script = ''
        if isinstance(content, str):
            full_script = content
        elif isinstance(content, dict):
            full_script = content.get('full_script', '')

            # 如果没有 full_script，尝试从 scenes 构建
            if not full_script:
                scenes = content.get('scenes', [])
                if isinstance(scenes, list) and len(scenes) > 0:
                    if isinstance(scenes[0], str):
                        full_script = '\n\n'.join(scenes)
                    elif isinstance(scenes[0], dict):
                        scene_texts = []
                        for scene in scenes:
                            scene_text = scene.get('text', '') or scene.get('content', '')
                            if scene_text:
                                scene_texts.append(scene_text)
                        full_script = '\n\n'.join(scene_texts)
        return full_script

    def _add_content_paragraphs(self, document: Document, full_script: str):
        """将剧本文本按段落添加到文档"""
        if full_script:
            paragraphs = full_script.split('\n')
            for para_text in paragraphs:
                if para_text.strip():
                    para = document.add_paragraph(para_text.strip())
                    for run in para.runs:
                        run.font.name = 'SimSun'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        else:
            document.add_paragraph('（暂无剧本内容）')

    def export_script(self, script_data: Dict[str, Any]) -> BytesIO:
        """导出单集剧本为Word文档"""
        document = self._setup_document()

        # 获取数据
        title = script_data.get('title', '未命名剧本')
        content = script_data.get('content', {})

        # 1. 剧本标题
        title_para = document.add_paragraph(title)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.bold = True
        title_run.font.size = Pt(24)
        title_run.font.name = 'SimHei'
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

        # 添加间距
        document.add_paragraph()

        # 2. 添加内容
        full_script = self._extract_content(content)
        self._add_content_paragraphs(document, full_script)

        # 保存到内存流
        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return buffer

    def export_merged(self, project_name: str, scripts_data: List[Dict[str, Any]]) -> BytesIO:
        """将多集剧本合并导出为一份Word文档"""
        document = self._setup_document()

        # 文档总标题
        main_title = document.add_paragraph(project_name)
        main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        main_run = main_title.runs[0]
        main_run.bold = True
        main_run.font.size = Pt(28)
        main_run.font.name = 'SimHei'
        main_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

        document.add_paragraph()

        # 遍历每集
        for i, script_data in enumerate(scripts_data):
            episode_number = script_data.get('episode_number', i + 1)
            title = script_data.get('title', f'第{episode_number}集')
            content = script_data.get('content', {})

            # 分集标题
            episode_title = document.add_paragraph(f'第{episode_number}集 {title}')
            episode_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
            ep_run = episode_title.runs[0]
            ep_run.bold = True
            ep_run.font.size = Pt(18)
            ep_run.font.name = 'SimHei'
            ep_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

            # 分隔线
            sep = document.add_paragraph('─' * 40)
            sep.runs[0].font.size = Pt(10)

            # 添加内容
            full_script = self._extract_content(content)
            self._add_content_paragraphs(document, full_script)

            # 集与集之间添加分页（最后一集不加）
            if i < len(scripts_data) - 1:
                document.add_page_break()

        # 保存到内存流
        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return buffer
